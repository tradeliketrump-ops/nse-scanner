"""
Generate VJs Structural Levels Pro V1 Documentation as Word (.docx)
Reads the markdown and produces a properly formatted Word document.
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re

def add_code_block(doc, code_text):
    """Add a code block with gray background styling."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    # Add shading to paragraph
    shading = p.paragraph_format.element.makeelement(qn('w:shd'), {
        qn('w:val'): 'clear',
        qn('w:color'): 'auto',
        qn('w:fill'): 'F0F0F5'
    })
    p.paragraph_format.element.get_or_add_pPr().append(shading)

def add_heading_styled(doc, text, level):
    """Add heading with consistent formatting."""
    h = doc.add_heading(text, level=level)
    return h

def add_table_from_rows(doc, headers, rows):
    """Add a formatted table with header row and data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(10)
    
    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)
    
    doc.add_paragraph()  # spacing after table

def generate():
    doc = Document()
    
    # ── Document Style ──
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    style.paragraph_format.space_after = Pt(6)
    
    # ── Title Page ──
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run('VJs Structural Levels Pro V1')
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = RGBColor(0x00, 0x8C, 0x6E)
    
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run('Complete Documentation')
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(40)
    run3 = p3.add_run('TradingView Pine Script Indicator\nNSE Cash Market — Scalp & Positional Trading')
    run3.font.size = Pt(12)
    run3.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
    
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(60)
    run4 = p4.add_run('Version 1.0 — June 2026')
    run4.font.size = Pt(11)
    run4.font.italic = True
    
    doc.add_page_break()
    
    # ── Table of Contents placeholder ──
    add_heading_styled(doc, 'Table of Contents', 1)
    toc_items = [
        '1. Overview',
        '2. Core Philosophy',
        '3. Section-by-Section Breakdown',
        '    3.1 Heikin Ashi Source',
        '    3.2 Previous Day Data & Pivot Calculation',
        '    3.3 Structural Levels',
        '    3.4 Trend Filters',
        '    3.5 Reversal Logic',
        '    3.6 Breakout / Breakdown Logic',
        '    3.7 One-Signal-Per-Trend Mechanism',
        '    3.8 RSI Extreme Exits',
        '    3.9 Moving Averages',
        '    3.10 Visual-Only Sections',
        '    3.11 Structural Level Lines',
        '    3.12 Right-Side Level Labels',
        '    3.13 Dashboard Table',
        '    3.14 Alert Conditions',
        '4. Signal Matrix Summary',
        '5. Trading Guidelines',
        '6. Parameter Reference',
        '7. Disclaimer',
    ]
    for item in toc_items:
        p = doc.add_paragraph(item)
        p.paragraph_format.space_after = Pt(2)
    
    doc.add_page_break()
    
    # ── 1. Overview ──
    add_heading_styled(doc, '1. Overview', 1)
    doc.add_paragraph(
        'VJs Structural Levels Pro V1 is a Pine Script indicator for TradingView that identifies '
        'high-probability intraday trade setups based on structural price levels derived from the '
        'previous day\'s price action. It combines daily pivot/ATR levels with Heikin Ashi smoothing, '
        'ADX/DMI trend filtering, and VWAP confirmation to generate actionable BUY and SELL signals '
        'across reversal and breakout/breakdown strategies.'
    )
    doc.add_paragraph(
        'The indicator is designed for NSE cash market trading and supports both scalp '
        '(ATR-based targets 1 & 2) and positional (fixed profit target) approaches within '
        'a single framework.'
    )
    
    # ── 2. Core Philosophy ──
    add_heading_styled(doc, '2. Core Philosophy', 1)
    doc.add_paragraph(
        'The indicator operates on a multi-timeframe confluence model:'
    )
    doc.add_paragraph(
        '1. Daily timeframe establishes the structural framework — pivot point, ATR-based support/resistance levels, '
        'and the broader trend context.',
        style='List Number'
    )
    doc.add_paragraph(
        '2. Intraday (current timeframe) executes the entry logic using Heikin Ashi candles for noise reduction '
        'and ADX/DMI for trend strength validation.',
        style='List Number'
    )
    doc.add_paragraph(
        '3. Signals are generated at the intersection of level-touch events and confirmed trend conditions.',
        style='List Number'
    )
    doc.add_paragraph(
        'No single component drives the signal — success comes from the confluence of structural levels, '
        'trend filter, and volatility context working together.'
    )
    
    # ── 3. Section-by-Section Breakdown ──
    add_heading_styled(doc, '3. Section-by-Section Breakdown', 1)
    
    # 3.1 Heikin Ashi
    add_heading_styled(doc, '3.1 Heikin Ashi Source', 2)
    add_code_block(doc, (
        'haTicker = ticker.heikinashi(syminfo.tickerid)\n'
        'haClose = request.security(haTicker, timeframe.period, close)\n'
        'haHigh  = request.security(haTicker, timeframe.period, high)\n'
        'haLow   = request.security(haTicker, timeframe.period, low)'
    ))
    doc.add_paragraph(
        'All level interactions and crossover logic use Heikin Ashi candles rather than raw price. '
        'Heikin Ashi filters out minor price noise and provides cleaner signals for level tests. '
        'HA close is used for EMA calculations, SMA calculations, crossover detection, and level-touch '
        'checks throughout the indicator.'
    )
    
    # 3.2 Previous Day Data
    add_heading_styled(doc, '3.2 Previous Day Data & Pivot Calculation', 2)
    add_code_block(doc, (
        'pdHigh  = request.security(syminfo.tickerid, "D", high[1], lookahead=barmerge.lookahead_on)\n'
        'pdLow   = request.security(syminfo.tickerid, "D", low[1], lookahead=barmerge.lookahead_on)\n'
        'pdClose = request.security(syminfo.tickerid, "D", close[1], lookahead=barmerge.lookahead_on)\n\n'
        'pivot = (pdHigh + pdLow + pdClose) / 3\n\n'
        'dailyATR = request.security(syminfo.tickerid, "D", ta.atr(14), lookahead=barmerge.lookahead_on)'
    ))
    doc.add_paragraph(
        'All daily values use the previous day\'s close ([1]) ensuring no intraday lookahead — the levels '
        'are fixed for the entire trading day based on data available before the market opened.'
    )
    doc.add_paragraph(
        'Pivot Formula: (Previous Day High + Previous Day Low + Previous Day Close) / 3. '
        'This is the classic floor-trader pivot point representing the market\'s "fair price" from the prior session.'
    )
    
    # 3.3 Structural Levels
    add_heading_styled(doc, '3.3 Structural Levels', 2)
    add_code_block(doc, (
        'sellRev  = pivot + dailyATR * sellRevATR      // pivot + ATR * 0.29\n'
        'buyRev   = pivot - dailyATR * buyRevATR       // pivot - ATR * 0.21\n\n'
        'breakout = pivot + dailyATR * breakoutATR     // pivot + ATR * 0.54\n'
        'breakdown= pivot - dailyATR * breakdownATR    // pivot - ATR * 0.46\n\n'
        'targetUp1 = breakout + dailyATR * target1ATR  // breakout + ATR * 0.28\n'
        'targetUp2 = breakout + dailyATR * target2ATR  // breakout + ATR * 0.41\n\n'
        'targetDn1 = breakdown - dailyATR * target1ATR // breakdown - ATR * 0.28\n'
        'targetDn2 = breakdown - dailyATR * target2ATR // breakdown - ATR * 0.46'
    ))
    
    add_table_from_rows(doc, 
        ['Level', 'Multiplier', 'Purpose'],
        [
            ['Target Up 2', 'Pivot + ATR \u00d7 0.82', 'Scalp profit target (aggressive)'],
            ['Target Up 1', 'Pivot + ATR \u00d7 0.69', 'Scalp profit target (conservative)'],
            ['Breakout', 'Pivot + ATR \u00d7 0.54', 'Bullish breakout trigger'],
            ['Sell Reversal', 'Pivot + ATR \u00d7 0.29', 'Bearish reversal zone'],
            ['Pivot', '(H+L+C)/3', 'Central reference / fair value'],
            ['Buy Reversal', 'Pivot \u2212 ATR \u00d7 0.21', 'Bullish reversal zone'],
            ['Breakdown', 'Pivot \u2212 ATR \u00d7 0.46', 'Bearish breakdown trigger'],
            ['Target Down 1', 'Breakdown \u2212 ATR \u00d7 0.28', 'Scalp profit target (conservative)'],
            ['Target Down 2', 'Breakdown \u2212 ATR \u00d7 0.46', 'Scalp profit target (aggressive)'],
        ]
    )
    doc.add_paragraph(
        'The multipliers are calibrated to capture typical intraday extension ranges observed in NSE stocks. '
        'The asymmetry between bullish and bearish multipliers accounts for the natural drift tendency of equities.'
    )
    
    # 3.4 Trend Filters
    add_heading_styled(doc, '3.4 Trend Filters', 2)
    doc.add_paragraph('EMA Cross (Fast/Slow):')
    add_code_block(doc, (
        'ema5  = ta.ema(haClose, emaFastLen)    // Default: 5\n'
        'ema20 = ta.ema(haClose, emaSlowLen)    // Default: 20'
    ))
    doc.add_paragraph('SMA 50:')
    add_code_block(doc, 'sma50 = ta.sma(haClose, smaLen)        // Default: 50')
    doc.add_paragraph('VWAP:')
    add_code_block(doc, 'vwapValue = ta.vwap(hlc3)')
    doc.add_paragraph('ADX/DMI:')
    add_code_block(doc, '[diplus, diminus, adx] = ta.dmi(adxLen, adxSmooth)')
    
    doc.add_paragraph('Complete Trend Conditions:')
    add_code_block(doc, (
        'bullTrend = ema5 > ema20 AND\n'
        '            close > sma50 AND\n'
        '            diplus > diminus AND\n'
        '            adx > adxThresh AND\n'
        '            bullVWAP (if enabled)'
    ))
    doc.add_paragraph(
        'Key Design Decision: All conditions must be true simultaneously. This prevents marginal setups '
        'from generating signals and ensures only high-confluence trades are considered.'
    )
    
    # 3.5 Reversal Logic
    add_heading_styled(doc, '3.5 Reversal Logic', 2)
    doc.add_paragraph('Touch Detection:')
    add_code_block(doc, (
        'buyTouch  = haClose > buyRev\n'
        'sellTouch = haClose < sellRev'
    ))
    doc.add_paragraph('Signal Confirmation (BUY-R):')
    add_code_block(doc, (
        'buyRevSignal = rsi < 80 AND haClose > buyRev AND bullTrend'
    ))
    doc.add_paragraph('BUY-R (Buy Reversal): ', style='List Bullet')
    doc.add_paragraph(
        'Triggered when Heikin Ashi close moves above the Buy Reversal level while the overall trend '
        'is bullish. This signals that a pullback to the reversal zone has found support, and price is '
        'resuming its uptrend. The RSI check (rsi < 80) prevents entering when the market is already overextended.'
    )
    doc.add_paragraph('Signal Confirmation (SELL-R):')
    add_code_block(doc, (
        'sellRevSignal = rsi > 20 AND haClose < sellRev AND bearTrend'
    ))
    doc.add_paragraph('SELL-R (Sell Reversal): ', style='List Bullet')
    doc.add_paragraph(
        'Triggered when HA close moves below the Sell Reversal level while the trend is bearish. '
        'Price has rallied into the resistance zone and reversed back down.'
    )
    
    # 3.6 Breakout
    add_heading_styled(doc, '3.6 Breakout / Breakdown Logic', 2)
    add_code_block(doc, (
        'buyBreakoutSignal  = crossover(haClose, breakout) AND bullTrend\n'
        'sellBreakdownSignal = crossunder(haClose, breakdown) AND bearTrend'
    ))
    doc.add_paragraph(
        'BUY-B (Buy Breakout): HA close crosses above the Breakout level — price is breaking out of its '
        'expected daily range with strong bullish momentum.'
    )
    doc.add_paragraph(
        'SELL-B (Sell Breakdown): HA close crosses below the Breakdown level — price is breaking down '
        'through support with strong bearish momentum.'
    )
    
    # 3.7 One-Signal-Per-Trend
    add_heading_styled(doc, '3.7 One-Signal-Per-Trend Mechanism', 2)
    add_code_block(doc, (
        'var int trendState = 0\n\n'
        'finalBuyR = buyRevSignal and trendState != 1\n'
        'finalSellR = sellRevSignal and trendState != -1\n\n'
        'if finalBuyR or finalBuyB\n'
        '    trendState := 1\n'
        '    targetPoints = profitTargetRs / tradeQty\n'
        '    longTargetPrice := close + targetPoints'
    ))
    doc.add_paragraph(
        'Prevents multiple signals in the same direction. Once a BUY signal fires, the trendState locks to +1 '
        'and no further BUY signals are generated until a SELL signal resets it. The fixed profit target is '
        'calculated as / (lotSize \u00d7 lots).'
    )
    
    # 3.8 RSI Exits
    add_heading_styled(doc, '3.8 RSI Extreme Exits', 2)
    add_code_block(doc, (
        'longRSIExit  = inLong AND crossover(rsi, 90)\n'
        'shortRSIExit = inShort AND crossunder(rsi, 10)'
    ))
    doc.add_paragraph(
        'Exits positions when RSI reaches statistical extremes (90+ for longs, 10 or below for shorts). '
        'These levels occur approximately 1-3% of the time and represent parabolic moves likely to reverse.'
    )
    
    # 3.9 MA
    add_heading_styled(doc, '3.9 Moving Averages', 2)
    add_code_block(doc, (
        'plot(ema5, "EMA 5", color=color.rgb(6, 121, 10), linewidth=2)\n'
        'plot(ema20, "EMA 20", color=color.red, linewidth=2)\n'
        'plot(sma50, "SMA 50", color=color.rgb(3, 3, 0))'
    ))
    doc.add_paragraph(
        'EMA5 (green), EMA20 (red), SMA50 (dark). Provide visual context for the trend filters.'
    )
    
    # 3.10 Visual-Only
    add_heading_styled(doc, '3.10 Visual-Only Sections', 2)
    doc.add_paragraph('The following sections are purely decorative and do not affect any signal logic:')
    doc.add_paragraph('HMA (Hull Moving Average)', style='List Bullet')
    doc.add_paragraph(
        'A visually smoothed moving average with color changes based on momentum. Rendered with a black border for emphasis.'
    )
    doc.add_paragraph('Coral Indicator', style='List Bullet')
    doc.add_paragraph(
        'A trend-following indicator based on cascading IIR filters. Available in standard, ribbon, and color bar modes. '
        'Includes HTF (1-hour) and LTF (10-minute) Coral states for additional context.'
    )
    doc.add_paragraph('HL/LH (Higher Low / Lower High)', style='List Bullet')
    doc.add_paragraph(
        'Detects swing structure patterns using adaptive pivot detection. Shows "HL" labels below bars during bullish '
        'trends and "LH" labels above bars during bearish trends.'
    )
    
    # 3.11 Level Lines
    add_heading_styled(doc, '3.11 Structural Level Lines', 2)
    doc.add_paragraph('At the start of each new day, the indicator draws:')
    doc.add_paragraph('Vertical anchor line (gray) — connects highest and lowest targets', style='List Bullet')
    doc.add_paragraph('Red dashed lines — scalp targets (Target 1 & Target 2)', style='List Bullet')
    doc.add_paragraph('Blue lines — Breakout (above, width=2) and Breakdown (below, width=2)', style='List Bullet')
    doc.add_paragraph('Red line — Sell Reversal level', style='List Bullet')
    doc.add_paragraph('Orange line — Pivot level (width=2)', style='List Bullet')
    doc.add_paragraph('Lime line — Buy Reversal level', style='List Bullet')
    doc.add_paragraph('All lines extend to the right, providing forward-looking reference levels.')
    
    # 3.12 Labels
    add_heading_styled(doc, '3.12 Right-Side Level Labels', 2)
    doc.add_paragraph(
        'When barstate.islast is true, price labels are drawn offset to the right. Each label shows '
        'the level name and current value (rounded to nearest integer). Example:'
    )
    add_code_block(doc, (
        'Target 2 - 5276\nTarget 1 - 5241\nBreakout - 5204\nSell Reversal - 5168\n'
        'Pivot - 5159\nBuy Reversal - 5152\nBreak Down - 5140\nTarget 1 - 5112\nTarget 2 - 5088'
    ))
    
    # 3.13 Dashboard
    add_heading_styled(doc, '3.13 Dashboard Table', 2)
    doc.add_paragraph('A compact 4-column table at the top center of the chart showing:')
    add_table_from_rows(doc,
        ['Column', 'Content', 'Text Color'],
        [
            ['ADX', 'Current ADX value', 'Yellow'],
            ['DI+', 'Positive directional index', 'Green'],
            ['DI-', 'Negative directional index', 'Red'],
            ['Status', 'BULL / BEAR / NEUTRAL', 'Green / Red / Orange'],
        ]
    )
    
    # 3.14 Alerts
    add_heading_styled(doc, '3.14 Alert Conditions', 2)
    add_table_from_rows(doc,
        ['Alert', 'Trigger', 'Use Case'],
        [
            ['BUY REVERSAL', 'BUY-R signal fires', 'Enter long on reversal'],
            ['SELL REVERSAL', 'SELL-R signal fires', 'Enter short on reversal'],
            ['BUY BREAKOUT', 'BUY-B signal fires', 'Enter long on breakout'],
            ['SELL BREAKDOWN', 'SELL-B signal fires', 'Enter short on breakdown'],
            ['LONG EXIT RSI', 'RSI > 90 while in long', 'Exit long at extreme'],
            ['SHORT EXIT RSI', 'RSI < 10 while in short', 'Exit short at extreme'],
        ]
    )
    
    # ── 4. Signal Matrix ──
    add_heading_styled(doc, '4. Signal Matrix Summary', 1)
    add_table_from_rows(doc,
        ['Signal', 'Level Touch', 'Trend Required', 'RSI Filter', 'Type'],
        [
            ['BUY-R', 'HA close > Buy Reversal', 'Bullish', 'RSI < 80', 'Reversal / pullback entry'],
            ['BUY-B', 'HA crossover Breakout', 'Bullish', 'None', 'Momentum breakout entry'],
            ['SELL-R', 'HA close < Sell Reversal', 'Bearish', 'RSI > 20', 'Reversal / pullback entry'],
            ['SELL-B', 'HA crossunder Breakdown', 'Bearish', 'None', 'Momentum breakdown entry'],
        ]
    )
    
    # ── 5. Trading Guidelines ──
    add_heading_styled(doc, '5. Trading Guidelines', 1)
    add_heading_styled(doc, 'Recommended Timeframes', 2)
    doc.add_paragraph('Primary: 1-hour or 30-minute charts', style='List Bullet')
    doc.add_paragraph('Minimum: 15-minute (requires at least 33 bars for ADX calculation)', style='List Bullet')
    doc.add_paragraph('Ideal setup: Use on 1H chart with Daily pivot levels', style='List Bullet')
    
    add_heading_styled(doc, 'Entry Rules', 2)
    doc.add_paragraph('Wait for a BUY-R / BUY-B or SELL-R / SELL-B label to appear', style='List Number')
    doc.add_paragraph('Verify the label corresponds with the current trend direction', style='List Number')
    doc.add_paragraph('Check the dashboard for ADX > 20 and correct DI+ / DI- alignment', style='List Number')
    doc.add_paragraph('For manual entry: enter on the next candle open after the signal', style='List Number')
    
    add_heading_styled(doc, 'Exit Rules', 2)
    doc.add_paragraph('Scalp Strategy:')
    doc.add_paragraph('Target 1: +0.28\u00d7 ATR from breakout (conservative)', style='List Bullet')
    doc.add_paragraph('Target 2: +0.41\u00d7 ATR from breakout (aggressive)', style='List Bullet')
    doc.add_paragraph('Stop: Breakout level for BUY-B, Breakdown level for SELL-B', style='List Bullet')
    doc.add_paragraph('Positional Strategy:')
    doc.add_paragraph('Fixed profit target (default: / position size)', style='List Bullet')
    doc.add_paragraph('5% trailing or fixed stop-loss', style='List Bullet')
    doc.add_paragraph('RSI extreme exits (RSI > 90 for longs, RSI < 10 for shorts)', style='List Bullet')
    
    add_heading_styled(doc, 'Risk Management', 2)
    doc.add_paragraph('Position size: lotSize \u00d7 lots (default: 65 \u00d7 10 = 650 shares)', style='List Bullet')
    doc.add_paragraph('Capital per position: 1.5L (built into the tracker system)', style='List Bullet')
    doc.add_paragraph('Maximum 1 active trade per direction (enforced by trendState)', style='List Bullet')
    
    # ── 6. Parameter Reference ──
    add_heading_styled(doc, '6. Parameter Reference', 1)
    add_table_from_rows(doc,
        ['Input', 'Default', 'Range', 'Effect'],
        [
            ['EMA Fast', '5', '3\u201320', 'Sensitivity of short-term trend'],
            ['EMA Slow', '20', '10\u201350', 'Medium-term trend direction'],
            ['SMA Length', '50', '20\u2013200', 'Major trend filter'],
            ['ADX Length', '14', '7\u201321', 'ADX lookback period'],
            ['ADX Smoothing', '14', '7\u201321', 'ADX smoothing'],
            ['ADX Threshold', '20', '15\u201330', 'Minimum trend strength'],
            ['Use VWAP', 'true', 'On/Off', 'Optional volume filter'],
            ['Sell Rev ATR Mult', '0.29', '0.1\u20130.5', 'Sell reversal distance from pivot'],
            ['Buy Rev ATR Mult', '0.21', '0.1\u20130.5', 'Buy reversal distance from pivot'],
            ['Breakout ATR Mult', '0.54', '0.3\u20130.8', 'Breakout distance from pivot'],
            ['Breakdown ATR Mult', '0.46', '0.3\u20130.8', 'Breakdown distance from pivot'],
            ['Target 1 ATR Mult', '0.28', '0.1\u20130.5', 'First scalp target distance'],
            ['Target 2 ATR Mult', '0.41', '0.1\u20130.6', 'Second scalp target distance'],
            ['Lot Size', '65', '1\u2013500', 'Lot/multiplier for position sizing'],
            ['Lots', '10', '1\u2013100', 'Number of lots'],
            ['Target ', '500000', '10000\u201310M', 'Fixed profit target in rupees'],
        ]
    )
    
    # ── 7. Disclaimer ──
    add_heading_styled(doc, '7. Disclaimer', 1)
    p = doc.add_paragraph()
    run = p.add_run(
        'This indicator is for educational and analytical purposes only. It does not constitute '
        'financial advice. Past performance (backtested or live) does not guarantee future results. '
        'Always practice proper risk management and consult a qualified financial advisor before '
        'making trading decisions.'
    )
    run.italic = True
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    # ── Save ──
    output_path = 'VJs_Structural_Levels_Pro_V1_Documentation.docx'
    doc.save(output_path)
    print(f'Document saved to {output_path}')

if __name__ == '__main__':
    generate()