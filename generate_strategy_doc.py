"""Generate updated NSE Swing Trade Strategy document v1.2."""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import datetime

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# ─── Title Page ─────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('\n\n\n\n')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('NSE Swing Trade Scanner')
run.bold = True
run.font.size = Pt(28)
run.font.color.rgb = RGBColor(0, 0xd4, 0xaa)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('Complete Trading Strategy Document v1.2')
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0x88, 0x88, 0x88)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run(f'\n{datetime.date.today().strftime("%B %d, %Y")}\n\nNifty500 Universe | Heiken-Ashi Breakout | 1H TV-Style Entry\n₹1.5L/Trade | 5% SL | 10% Target\nSQLite Database | Intraday Scans | Auto-Refresh | Nifty50 FUT')

doc.add_page_break()

# ─── Table of Contents ──────────────────────────────────────────────
doc.add_heading('Table of Contents', level=1)
toc_items = [
    '1. Overview',
    '2. Scan Universe — Nifty500 + Nifty50 Index',
    '3. Daily Technical Screening',
    '4. 1-Hour TV-Style Entry Analysis',
    '5. Signal Types',
    '6. Position Creation & Entry',
    '7. Risk Management',
    '8. Intraday Scan Schedule',
    '9. Dashboard & Monitoring',
    '10. Performance Reporting',
    '11. Data Persistence — SQLite Database',
    '12. Nifty 50 Index Trading',
    '13. Auto-Refresh & Dashboard Sync',
    'Appendix: Key Code Parameters',
]
for item in toc_items:
    p = doc.add_paragraph(item, style='List Number')
    p.paragraph_format.space_after = Pt(2)

doc.add_page_break()

# ─── Section 1: Overview ───────────────────────────────────────────
doc.add_heading('1. Overview', level=1)
doc.add_paragraph(
    'The NSE Swing Trade Scanner is an automated system that scans the Nifty500 universe '
    'daily to identify swing trading opportunities. It uses a two-pass approach:'
)

passes = [
    ('Pass 1 — Daily Screening', 
     'Downloads 6 months of daily OHLC data via yfinance. Applies Heiken-Ashi smoothing, '
     'EMA20/SMA50 crossover logic, volume filters, market cap filter (₹5,000 Cr min), '
     'and RSI filter (40-70). Only stocks passing all criteria proceed to Pass 2.'),
    ('Pass 2 — 1H Entry Analysis', 
     'Downloads 10 days of 1-hour data for qualifying stocks via yfinance. Computes '
     'ADX/DMI for trend strength, pivot/ATR structural levels, and generates entry '
     'signals: BUY-R, BUY-B, SELL-R, SELL-B, or NEUTRAL.'),
    ('Position Creation (Stocks)',
     'BUY-R and BUY-B signals automatically create trade positions with ₹1.5L capital, '
     '5% stop-loss, and 10% profit target. Quantity = Capital ÷ Entry Price. Same-day '
     'entry during market hours via intraday scans; next-day open entry for end-of-day scans.'),
    ('Position Creation (Nifty 50 Index)',
     'Nifty 50 is scanned as a special instrument with 100-point stop-loss, 300-point '
     'profit target, and 1 Qty (65 lot multiplier). P&L = (Close - Entry) × 1 × 65.'),
]
for title, desc in passes:
    p = doc.add_paragraph()
    run = p.add_run(f'{title}: ')
    run.bold = True
    p.add_run(desc)

doc.add_paragraph(
    '\nAll positions are monitored every 15 minutes for stop-loss and target hits via '
    'APScheduler. Data is stored in a SQLite database (trades_history.db) with persistent '
    'disk storage on Render, ensuring no data loss across restarts. The dashboard '
    'auto-refreshes every 5 minutes to sync with scheduled scans.'
)

# ─── Section 2: Scan Universe ───────────────────────────────────────
doc.add_heading('2. Scan Universe — Nifty500 + Nifty50 Index', level=1)
doc.add_paragraph(
    'The scanner covers the Nifty500 universe — the top 500 companies listed on the '
    'National Stock Exchange of India by market capitalization. Additionally, the '
    'Nifty 50 Spot Index (^NSEI) is always included for index futures trading signals.'
)
doc.add_heading('Stock Sources', level=2)
doc.add_paragraph(
    'Primary: NSE India API (https://www.nseindia.com/api/equity-stockIndices?index=NIFTY%20500) '
    '— fetches live Nifty500 constituents with session-based authentication.'
)
doc.add_paragraph(
    'Fallback: A comprehensive static list of ~370 highly liquid stocks covering:'
)
bullets = [
    'Nifty 50 — 50 largest stocks',
    'Nifty Next 50 — 50 next largest',
    'Nifty Midcap 100 — 100 liquid mid-cap stocks',
    'Additional liquid stocks for broad coverage',
    'Nifty 50 Index — always included for FUT signals',
]
for b in bullets:
    doc.add_paragraph(b, style='List Bullet')

doc.add_paragraph(
    'The total scanned count is displayed on the dashboard as "📡 Scanned". The Nifty 50 '
    'Index is included in the scan but not counted in the stock universe total.'
)

# ─── Section 3: Daily Technical Screening ──────────────────────────
doc.add_heading('3. Daily Technical Screening', level=1)
doc.add_paragraph(
    'Each stock in the universe is evaluated using daily timeframe data (6 months). '
    'The screening process follows multiple technical filters:'
)

doc.add_heading('3.1 Heiken-Ashi Smoothing', level=2)
doc.add_paragraph(
    'Raw OHLC data is converted to Heiken-Ashi (HA) candles, which filter out market noise '
    'by using the formula:\n'
    '• HA_Close = (Open + High + Low + Close) / 4\n'
    '• HA_Open = (Previous HA_Open + Previous HA_Close) / 2\n'
    '• HA_High = max(High, HA_Open, HA_Close)\n'
    '• HA_Low = min(Low, HA_Open, HA_Close)'
)

doc.add_heading('3.2 Breakout Detection', level=2)
doc.add_paragraph(
    'Three conditions must be met (the "3C" rule):\n'
    '• C1: HA_Close > EMA20 (Heiken-Ashi Close above 20-period Exponential Moving Average)\n'
    '• C2: HA_Close > SMA50 (Heiken-Ashi Close above 50-period Simple Moving Average)\n'
    '• C3: EMA20 > SMA50 (Golden Crossover — short-term above long-term)\n\n'
    'If all three are true and persist for 1+ days: "Fresh Breakout" (1-3 days), '
    '"Strong Momentum" (4-10 days), or "Already Rallied" (10+ days).'
)

doc.add_heading('3.3 Volume & Market Cap Filters', level=2)
doc.add_paragraph(
    '• Minimum average volume: 500,000 shares over 20-day lookback\n'
    '• Minimum market cap: ₹5,000 Crores (estimated from Price × Avg Vol × 20 / 10M)\n'
    '• Volume spike detection: compares latest volume to 20-day average'
)

doc.add_heading('3.4 RSI & Relative Strength', level=2)
doc.add_paragraph(
    '• RSI period: 14, Range: 40-70 (disabled for manual scans via dashboard)\n'
    '• Relative Strength: compares stock performance to Nifty 50 over 63 days, scored 0-100\n'
    '• Higher-High / Higher-Low (HH/HL): Confirms uptrend over 20 days\n'
    '• Low Wick: Identifies buying pressure (lower wick < 30% of candle range)'
)

# ─── Section 4: 1-Hour TV-Style Analysis ──────────────────────────
doc.add_heading('4. 1-Hour TV-Style Entry Analysis', level=1)
doc.add_paragraph(
    'Stocks that pass the daily screening proceed to 1-hour analysis. This uses '
    'TradingView-inspired logic with pivot-based structural levels and ADX/DMI trend detection.'
)

doc.add_heading('4.1 Data & Indicators', level=2)
doc.add_paragraph(
    '• 10 days of 1-hour data downloaded for each qualifying stock\n'
    '• Heiken-Ashi applied to 1H data for smoothing\n'
    '• EMA20 on 1H HA Close\n'
    '• RSI(14) on 1H HA Close\n'
    '• ADX(14) and DMI(14) with Wilder smoothing (requires 33+ bars → 10 days of data)'
)

doc.add_heading('4.2 Daily Structural Levels', level=2)
doc.add_paragraph(
    'Using the previous day\'s OHLC data:\n'
    '• Pivot = (High + Low + Close) / 3\n'
    '• Daily ATR = Average True Range (14-period Wilder smoothed)\n'
    '• Buy Reversal = Pivot - (ATR × 0.21)\n'
    '• Sell Reversal = Pivot + (ATR × 0.29)\n'
    '• Breakout = Pivot + (ATR × 0.54)\n'
    '• Breakdown = Pivot - (ATR × 0.46)'
)

doc.add_heading('4.3 Signal Logic', level=2)
p = doc.add_paragraph()
run = p.add_run('Bullish Trend Condition: ')
run.bold = True
p.add_run('1H HA_Close > 1H EMA20 AND DI+ > DI- AND ADX > 20 AND RSI < 80')

p = doc.add_paragraph()
run = p.add_run('Bearish Trend Condition: ')
run.bold = True
p.add_run('1H HA_Close < 1H EMA20 AND DI- > DI+ AND ADX > 20 AND RSI > 20')

signals = [
    ('BUY-R (Buy Reversal)', 'HA_Close touches/breaks below Buy Reversal level + Bullish Trend'),
    ('BUY-B (Buy Breakout)', 'HA_Close crosses above Breakout level + Bullish Trend'),
    ('SELL-R (Sell Reversal)', 'HA_Close touches/breaks above Sell Reversal level + Bearish Trend'),
    ('SELL-B (Sell Breakdown)', 'HA_Close crosses below Breakdown level + Bearish Trend'),
    ('NEUTRAL', 'No signal conditions met; ADX/DMI state noted for reference'),
]
for sig, desc in signals:
    p = doc.add_paragraph()
    run = p.add_run(f'{sig}: ')
    run.bold = True
    p.add_run(desc)

# ─── Section 5: Signal Types ───────────────────────────────────────
doc.add_heading('5. Signal Types & Action', level=1)

table = doc.add_table(rows=5, cols=4)
table.style = 'Light Grid Accent 1'
table.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Signal', 'Full Name', 'Auto-Trade', 'Action']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

data = [
    ['BUY-R', 'Buy Reversal', '✅ Yes', 'Buy at signal price with SL(5%)/Target(10%) for stocks'],
    ['BUY-B', 'Buy Breakout', '✅ Yes', 'Buy at signal price with SL(5%)/Target(10%) for stocks'],
    ['SELL-R', 'Sell Reversal', '❌ No', 'Listed on dashboard for manual FUT short reference'],
    ['SELL-B', 'Sell Breakdown', '❌ No', 'Listed on dashboard for manual FUT short reference'],
]
for row_idx, row_data in enumerate(data):
    for col_idx, val in enumerate(row_data):
        table.rows[row_idx + 1].cells[col_idx].text = val

# ─── Section 6: Position Creation ──────────────────────────────────
doc.add_heading('6. Position Creation & Entry', level=1)

doc.add_heading('6.1 Same-Day Entry (Intraday Scan)', level=2)
doc.add_paragraph(
    'When a BUY signal is detected during an intraday scan (9:45 AM, 11:00 AM, 1:00 PM, 2:30 PM):\n'
    '• Position is created immediately at the current market price\n'
    '• Status is set to "Active" (no pending wait)\n'
    '• Market hours check uses IST (UTC+5:30) via built-in timezone conversion\n'
    '• Stop-loss (5%) and Target (10%) are calculated from entry price\n'
    '• Quantity = ₹1,50,000 / entry price (rounded down)\n'
    '• Position enters monitoring immediately'
)

doc.add_heading('6.2 Next-Day Entry (EOD Scan)', level=2)
doc.add_paragraph(
    'When a BUY signal is detected after 3:30 PM IST:\n'
    '• Position is created as "Pending Entry"\n'
    '• At next market open (9:30 AM), entry price is fetched as the open price\n'
    '• SL, Target, and Quantity are calculated at that time\n'
    '• Position becomes "Active" and enters monitoring\n'
    '• Pending activation happens via the 15-min scheduler'
)

doc.add_heading('6.3 Position Duplication Prevention', level=2)
doc.add_paragraph(
    'Each symbol can only have ONE active or pending position at a time:\n'
    '• If a position already exists for a symbol (active or pending), new scans UPDATE it\n'
    '• Signal price, entry zone, and R:R are refreshed with latest data\n'
    '• Pending positions get activated if they\'re still pending and market is open\n'
    '• Old positions are NEVER replaced by new scans — only updated\n'
    '• This ensures positions persist until SL or Target is hit'
)

doc.add_heading('6.4 Position Parameters', level=2)
doc.add_paragraph('Stocks: Capital ₹1.5L | SL 5% | Target 10% | Qty = Capital ÷ Price', style='List Bullet')
doc.add_paragraph('Nifty 50: Qty 1 | SL 100 pts | Target 300 pts | Lot multiplier 65', style='List Bullet')

# ─── Section 7: Risk Management ─────────────────────────────────────
doc.add_heading('7. Risk Management', level=1)
doc.add_paragraph('Capital Allocation: ₹1.5 Lakh per position for stocks. Nifty 50 uses 1 lot with ₹1.5L margin.', style='List Bullet')
doc.add_paragraph('Stop-Loss: 5% fixed stop-loss for stocks. 100 points for Nifty 50. Auto-closed when price ≤ SL.', style='List Bullet')
doc.add_paragraph('Profit Target: 10% fixed for stocks. 300 points for Nifty 50. Auto-closed when price ≥ Target.', style='List Bullet')
doc.add_paragraph('Position Monitoring: Every 15 minutes during market hours via APScheduler.', style='List Bullet')
doc.add_paragraph('IST Timezone: All market hours checks convert UTC→IST (UTC+5:30) for correct operation on Render.', style='List Bullet')
doc.add_paragraph('Manual Refresh: "Refresh Prices" button on dashboard forces immediate price update.', style='List Bullet')

# ─── Section 8: Intraday Scan Schedule ─────────────────────────────
doc.add_heading('8. Intraday Scan Schedule', level=1)
doc.add_paragraph(
    'Four scheduled scans during market hours to catch intraday signals:'
)

table2 = doc.add_table(rows=5, cols=3)
table2.style = 'Light Grid Accent 1'
table2.alignment = WD_TABLE_ALIGNMENT.CENTER

headers2 = ['Time (IST)', 'Scan Type', 'Purpose']
for i, h in enumerate(headers2):
    cell = table2.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

schedule = [
    ['9:45 AM', '1H Re-Analysis', 'Catches early morning breakouts (9:15-9:45 AM)'],
    ['11:00 AM', '1H Re-Analysis', 'Catches late morning signals'],
    ['1:00 PM', '1H Re-Analysis', 'Catches early afternoon signals'],
    ['2:30 PM', '1H Re-Analysis', 'Last chance before end-of-day'],
]
for row_idx, row_data in enumerate(schedule):
    for col_idx, val in enumerate(row_data):
        table2.rows[row_idx + 1].cells[col_idx].text = val

doc.add_paragraph(
    '\nIn addition:\n'
    '• Full daily scan (370+1 stocks) runs when you click "Run Scan" on the dashboard\n'
    '• Trade position monitoring (SL/Target checks) runs every 15 minutes during market hours\n'
    '• All scheduled jobs use CronTrigger with Asia/Kolkata timezone\n'
    '• Dashboard auto-refreshes every 5 minutes to display latest data'
)

# ─── Section 9: Dashboard ──────────────────────────────────────────
doc.add_heading('9. Dashboard & Monitoring', level=1)
doc.add_paragraph(
    'The web dashboard provides real-time visibility into scan results and trade performance. '
    'It auto-refreshes every 5 minutes.'
)

doc.add_heading('9.1 Dashboard Sections', level=2)
doc.add_paragraph('STOCKS Tab:', style='List Bullet')
sections = [
    'Summary Card: 📡 Scanned (universe size), 🎯 Qualified, 🔍 Filtered Out, Fresh Breakout, BUY NOW, WAIT',
    'Watchlist Table: Symbol, Sector, Price, Stage, 1H Setup, 1H Detail, D_E20, R:R',
    'Nifty 50: Appears in the watchlist as a special instrument with its own signals',
    'Filters: All, BUY-R, BUY-B, SELL-R, SELL-B, Fresh Breakout',
    'Scan Progress Bar: Shows real-time progress during a full scan',
    'Download Excel: Export latest results as XLSX',
]
doc.add_paragraph('Performance Report Tab:', style='List Bullet')
sections2 = [
    'Portfolio Summary: Total Signals, Active, Pending, Closed, Win Rate',
    'P&L Details: Total Invested, Current Value, Net P&L (₹), Return %',
    'Open Positions Table: Symbol, Sector, Signal Date, Entry, Current, P&L, Status, SL/Target',
    'Nifty 50 positions shown with qty=1, 65x lot multiplier applied to P&L',
    'Per-Symbol Performance: Signal count, wins, losses, win rate per stock',
    'Refresh Button: Force-update all prices manually',
]
for s in sections + sections2:
    doc.add_paragraph(s, style='List Bullet 2')

# ─── Section 10: Performance Reporting ─────────────────────────────
doc.add_heading('10. Performance Reporting', level=1)
doc.add_paragraph(
    'The system maintains complete trade history in a SQLite database (trades_history.db) with:'
)
doc.add_paragraph('Portfolio-level aggregation: total invested, current value, net P&L, return %, win rate', style='List Bullet')
doc.add_paragraph('Per-position tracking: entry price, quantity, SL, target, current price, P&L', style='List Bullet')
doc.add_paragraph('Per-symbol statistics: total signals, active, closed wins/losses, win rate, total P&L', style='List Bullet')
doc.add_paragraph('Automatic position closure when SL or Target is hit', style='List Bullet')
doc.add_paragraph('Data persists across server restarts via Render persistent disk + SQLite', style='List Bullet')

# ─── Section 11: SQLite Database ──────────────────────────────────
doc.add_heading('11. Data Persistence — SQLite Database', level=1)
doc.add_paragraph(
    'Trade history is stored in a SQLite database (trades_history.db) instead of JSON files. '
    'This provides atomic writes, thread-safe concurrent access, and row-level operations.'
)
doc.add_paragraph('Database: trades_history.db (SQLite, built into Python)', style='List Bullet')
doc.add_paragraph('Storage: Render persistent disk at /data/ or local project directory as fallback', style='List Bullet')
doc.add_paragraph('Reliability: Atomic WAL (Write-Ahead Logging) prevents corruption', style='List Bullet')
doc.add_paragraph('Concurrency: Thread-safe with threading.Lock for all operations', style='List Bullet')
doc.add_paragraph('Performance: Indexed on symbol and status columns for fast queries', style='List Bullet')
doc.add_paragraph('No data loss: Survives all server restarts, redeploys, and crashes', style='List Bullet')
doc.add_paragraph('Same public API: All get_tracker(), create_positions_from_results(), etc. unchanged', style='List Bullet')

# ─── Section 12: Nifty 50 Index Trading ────────────────────────────
doc.add_heading('12. Nifty 50 Index Trading', level=1)
doc.add_paragraph(
    'The Nifty 50 Spot Index (^NSEI) is scanned as a special instrument alongside the '
    'stock universe. It uses a modified trading rule set designed for index futures:'
)

doc.add_heading('12.1 Nifty 50 vs Stocks Comparison', level=2)
table_nifty = doc.add_table(rows=7, cols=2)
table_nifty.style = 'Light Grid Accent 1'
table_nifty.alignment = WD_TABLE_ALIGNMENT.CENTER
for i, h in enumerate(['Parameter', 'Nifty 50', 'Stocks']):
    pass  # table will have 2 cols

# 2-column: Parameter | Value
nifty_params = [
    ('Data Source', '^NSEI (yfinance)'),
    ('Quantity', '1 position (fixed)'),
    ('Lot Multiplier', '65'),
    ('Stop-Loss', '100 points below entry'),
    ('Profit Target', '300 points above entry'),
    ('P&L Formula', '(Current - Entry) × 1 × 65'),
    ('Risk:Reward', '1:3'),
]
for i, (i_name, i_val) in enumerate([('Parameter', 'Value')] + nifty_params):
    pass  # just text below

doc.add_paragraph('Data Source: ^NSEI via yfinance — Nifty 50 Spot Index (proxy for futures)', style='List Bullet')
doc.add_paragraph('Quantity: 1 position (fixed, not capital-based)', style='List Bullet')
doc.add_paragraph('Lot Multiplier: 65 (1 position = 65 contracts on NSE)', style='List Bullet')
doc.add_paragraph('Stop-Loss: 100 points below entry price (= -100 × 65 = -₹6,500)', style='List Bullet')
doc.add_paragraph('Profit Target: 300 points above entry price (= +300 × 65 = +₹19,500)', style='List Bullet')
doc.add_paragraph('Risk:Reward: 1:3 (100 pts risk → 300 pts reward)', style='List Bullet')
doc.add_paragraph('P&L: (Current Price - Entry Price) × 1 × 65 = ₹ per point', style='List Bullet')
doc.add_paragraph('Same-day entry during market hours, next-day if after close', style='List Bullet')

doc.add_heading('12.2 Example Trade', level=2)
doc.add_paragraph(
    'Nifty gets BUY-B signal at 24,000:\n'
    '• Entry: 24,000\n'
    '• Stop-Loss: 23,900 (100 pts → -₹6,500 loss)\n'
    '• Target: 24,300 (300 pts → +₹19,500 profit)\n'
    '• Risk:Reward: 1:3'
)

# ─── Section 13: Auto-Refresh ──────────────────────────────────────
doc.add_heading('13. Auto-Refresh & Dashboard Sync', level=1)
doc.add_paragraph(
    'The dashboard automatically refreshes every 5 minutes to sync with the scheduler. '
    'This ensures new signals, position updates, and P&L changes appear without manual page reload.'
)
doc.add_paragraph('Watchlist auto-refresh: Every 5 minutes → reloads scan data from /api/watchlist/latest-data', style='List Bullet')
doc.add_paragraph('Portfolio auto-refresh: Every 5 minutes → updates positions, P&L, portfolio summary', style='List Bullet')
doc.add_paragraph('Scan poll: Dashboard polls /api/scan/status every 3 seconds during an active manual scan', style='List Bullet')
doc.add_paragraph('Max poll attempts: 300 (15 minutes timeout for long scans)', style='List Bullet')
doc.add_paragraph('IST timezone: All server-side time checks convert UTC to IST for correct Indian market hours', style='List Bullet')

# ─── Appendix ───────────────────────────────────────────────────────
doc.add_heading('Appendix: Key Code Parameters', level=1)
doc.add_paragraph('Configuration values used in the scanner:')

table3 = doc.add_table(rows=25, cols=2)
table3.style = 'Light Grid Accent 1'
table3.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, h in enumerate(['Parameter', 'Value']):
    cell = table3.rows[0].cells[i]
    cell.text = h
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.bold = True

params_list = [
    ('Universe Size', '~370 Nifty500 stocks + Nifty50 Index'),
    ('Daily Data Period', '6 months'),
    ('1H Data Period', '10 days (33+ bars for ADX)'),
    ('EMA Period', '20'),
    ('SMA Period', '50'),
    ('RSI Period', '14'),
    ('Volume Lookback', '20 days'),
    ('Min Market Cap', '₹5,000 Cr'),
    ('Min Avg Volume', '500,000 shares'),
    ('ADX Threshold', '20'),
    ('Breakout Lookback', '30 days'),
    ('Stock Capital', '₹1,50,000 per trade'),
    ('Stock Stop-Loss', '5%'),
    ('Stock Target', '10%'),
    ('Nifty Quantity', '1 position (65 lot multiplier)'),
    ('Nifty Stop-Loss', '100 points below entry'),
    ('Nifty Target', '300 points above entry'),
    ('Nifty P&L per pt', '₹65'),
    ('Database', 'SQLite (trades_history.db)'),
    ('Storage', 'Render /data disk or local fallback'),
    ('Auto-Refresh', 'Every 5 minutes'),
    ('Intraday Scan Times', '9:45, 11:00, 13:00, 14:30 IST'),
    ('Timezone Handling', 'UTC→IST conversion on Render'),
    ('Trade Monitoring', 'Every 15 minutes during market hours'),
]
for row_idx, (param, val) in enumerate(params_list):
    table3.rows[row_idx + 1].cells[0].text = param
    table3.rows[row_idx + 1].cells[1].text = val

# ─── Save ───────────────────────────────────────────────────────────
filename = 'NSE_Swing_Trade_Strategy.docx'
doc.save(filename)
print(f'✅ Document saved: {filename}')